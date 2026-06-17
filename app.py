"""抖音脚本生成器 - Flask 主应用"""
import os
import json
from flask import Flask, render_template, request, jsonify, send_from_directory
from config import Config
from models import db, Plan, KnowledgeDoc, KnowledgeCard, Video, Script
from utils.deepseek import generate_script, generate_plan, analyze_video_content
from utils.document_parser import parse_document, split_into_cards
from utils.video_parser import parse_video_url
from datetime import datetime, timezone

def create_app():
    app = Flask(__name__)
    app.config.from_object(Config)
    db.init_app(app)

    # 确保上传目录存在
    os.makedirs(Config.UPLOAD_FOLDER, exist_ok=True)
    # SQLite 需要确保数据库文件目录存在
    if 'sqlite' in Config.SQLALCHEMY_DATABASE_URI:
        db_path = Config.SQLALCHEMY_DATABASE_URI.replace('sqlite:///', '')
        os.makedirs(os.path.dirname(db_path), exist_ok=True)

    with app.app_context():
        db.create_all()

    # ==================== 页面路由 ====================

    @app.route('/')
    @app.route('/home')
    def home():
        return render_template('home.html')

    @app.route('/planning')
    def planning():
        return render_template('planning.html')

    @app.route('/knowledge')
    def knowledge():
        return render_template('knowledge.html')

    @app.route('/accounts')
    def accounts():
        return render_template('accounts.html')

    @app.route('/history')
    def history():
        return render_template('history.html')

    # ==================== API: 脚本生成 ====================

    @app.route('/api/generate', methods=['POST'])
    def api_generate():
        data = request.get_json()
        api_key = data.get('api_key', '')
        title = data.get('title', '')
        requirement = data.get('requirement', '')
        creative_ratio = data.get('creative_ratio', 50)
        plan_ids = data.get('plan_ids', [])
        knowledge_ids = data.get('knowledge_ids', [])
        video_ids = data.get('video_ids', [])

        if not title or not requirement:
            return jsonify({"error": "请填写脚本标题和创作需求"}), 400

        if not api_key:
            return jsonify({"error": "请先设置 DeepSeek API Key"}), 400

        # 获取参考内容
        plan_refs = [Plan.query.get(pid).to_dict() for pid in plan_ids if Plan.query.get(pid)]
        knowledge_refs = [KnowledgeCard.query.get(kid).to_dict() for kid in knowledge_ids if KnowledgeCard.query.get(kid)]
        video_refs = [Video.query.get(vid).to_dict() for vid in video_ids if Video.query.get(vid)]
        fast_mode = data.get('fast_mode', False)

        try:
            result = generate_script(api_key, title, requirement, creative_ratio,
                                     plan_refs, knowledge_refs, video_refs,
                                     fast_mode=fast_mode)

            # 提取关键词（简单的标签提取）
            keywords = extract_keywords(requirement)

            # 保存脚本
            script = Script(
                title=title,
                content=result,
                requirement=requirement,
                creative_ratio=creative_ratio,
                plan_refs=json.dumps(plan_ids),
                knowledge_refs=json.dumps(knowledge_ids),
                video_refs=json.dumps(video_ids),
                keywords=keywords,
            )
            db.session.add(script)
            db.session.commit()

            return jsonify({"success": True, "script": script.to_dict()})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route('/api/scripts', methods=['GET', 'POST'])
    def api_scripts():
        if request.method == 'POST':
            data = request.get_json()
            script = Script(
                title=data.get('title', ''),
                content=data.get('content', ''),
                requirement=data.get('requirement', ''),
                keywords=extract_keywords(data.get('content', '')[:200]),
            )
            db.session.add(script)
            db.session.commit()
            return jsonify({"success": True, "script": script.to_dict()})

        scripts = Script.query.order_by(Script.created_at.desc()).all()
        return jsonify([s.to_dict() for s in scripts])

    @app.route('/api/scripts/<int:script_id>', methods=['GET', 'DELETE'])
    def api_script(script_id):
        script = Script.query.get_or_404(script_id)
        if request.method == 'DELETE':
            db.session.delete(script)
            db.session.commit()
            return jsonify({"success": True})
        return jsonify(script.to_dict())

    # ==================== API: 前期策划 ====================

    @app.route('/api/plans', methods=['GET', 'POST'])
    def api_plans():
        if request.method == 'POST':
            data = request.get_json()
            plan = Plan(
                title=data.get('title', ''),
                content=data.get('content', ''),
                source_type=data.get('source_type', 'manual'),
            )
            db.session.add(plan)
            db.session.commit()
            return jsonify({"success": True, "plan": plan.to_dict()})

        plans = Plan.query.order_by(Plan.updated_at.desc()).all()
        return jsonify([p.to_dict() for p in plans])

    @app.route('/api/plans/<int:plan_id>', methods=['GET', 'PUT', 'DELETE'])
    def api_plan(plan_id):
        plan = Plan.query.get_or_404(plan_id)
        if request.method == 'DELETE':
            db.session.delete(plan)
            db.session.commit()
            return jsonify({"success": True})
        elif request.method == 'PUT':
            data = request.get_json()
            plan.title = data.get('title', plan.title)
            plan.content = data.get('content', plan.content)
            db.session.commit()
            return jsonify({"success": True, "plan": plan.to_dict()})
        return jsonify(plan.to_dict())

    @app.route('/api/planning/chat', methods=['POST'])
    def api_planning_chat():
        """AI 聊天式策划"""
        data = request.get_json()
        api_key = data.get('api_key', '')
        history = data.get('history', [])

        if not api_key:
            return jsonify({"error": "请先设置 DeepSeek API Key"}), 400

        try:
            from utils.deepseek import chat_planning
            result = chat_planning(api_key, history)
            return jsonify(result)
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route('/api/plans/ai-create', methods=['POST'])
    def api_plan_ai():
        data = request.get_json()
        api_key = data.get('api_key', '')
        topic = data.get('topic', '')
        requirements = data.get('requirements', '')

        if not topic:
            return jsonify({"error": "请输入策划主题"}), 400
        if not api_key:
            return jsonify({"error": "请先设置 DeepSeek API Key"}), 400

        try:
            result = generate_plan(api_key, topic, requirements)
            plan = Plan(
                title=topic,
                content=result,
                source_type='ai',
            )
            db.session.add(plan)
            db.session.commit()
            return jsonify({"success": True, "plan": plan.to_dict()})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route('/api/plans/upload', methods=['POST'])
    def api_plan_upload():
        file = request.files.get('file')
        if not file:
            return jsonify({"error": "请选择文件"}), 400

        filename = file.filename
        ext = os.path.splitext(filename)[1].lower()
        if ext not in {'.pdf', '.doc', '.docx', '.txt', '.md'}:
            return jsonify({"error": "不支持的文件格式，请上传 PDF/Word/TXT/MD"}), 400

        filepath = os.path.join(Config.UPLOAD_FOLDER, filename)
        file.save(filepath)

        try:
            content = parse_document(filepath)
            plan = Plan(
                title=os.path.splitext(filename)[0],
                content=content,
                source_type='upload',
                file_path=filepath,
            )
            db.session.add(plan)
            db.session.commit()
            return jsonify({"success": True, "plan": plan.to_dict()})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    # ==================== API: 知识库 ====================

    @app.route('/api/knowledge/upload', methods=['POST'])
    def api_knowledge_upload():
        file = request.files.get('file')
        if not file:
            return jsonify({"error": "请选择文件"}), 400

        filename = file.filename
        ext = os.path.splitext(filename)[1].lower()
        if ext not in {'.pdf', '.doc', '.docx', '.txt', '.md'}:
            return jsonify({"error": "不支持的文件格式"}), 400

        filepath = os.path.join(Config.UPLOAD_FOLDER, filename)
        file.save(filepath)

        try:
            content = parse_document(filepath)

            # 创建知识文档
            doc = KnowledgeDoc(
                title=os.path.splitext(filename)[0],
                file_name=filename,
                file_path=filepath,
                content=content,
            )
            db.session.add(doc)
            db.session.flush()

            # 自动拆解为知识卡片
            cards_data = split_into_cards(content)
            for card_data in cards_data:
                card = KnowledgeCard(
                    doc_id=doc.id,
                    title=card_data['title'],
                    content=card_data['content'],
                    tags=card_data['tags'],
                )
                db.session.add(card)

            doc.card_count = len(cards_data)
            db.session.commit()

            cards = KnowledgeCard.query.filter_by(doc_id=doc.id).all()
            return jsonify({
                "success": True,
                "doc": doc.to_dict(),
                "cards": [c.to_dict() for c in cards],
            })
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route('/api/knowledge/docs', methods=['GET', 'DELETE'])
    def api_knowledge_docs():
        if request.method == 'DELETE':
            doc_id = request.args.get('id')
            if doc_id:
                doc = KnowledgeDoc.query.get_or_404(doc_id)
                KnowledgeCard.query.filter_by(doc_id=doc.id).delete()
                db.session.delete(doc)
                db.session.commit()
                return jsonify({"success": True})
            return jsonify({"error": "缺少文档ID"}), 400

        docs = KnowledgeDoc.query.order_by(KnowledgeDoc.created_at.desc()).all()
        return jsonify([d.to_dict() for d in docs])

    @app.route('/api/knowledge/cards', methods=['GET'])
    def api_knowledge_cards():
        doc_id = request.args.get('doc_id')
        if doc_id:
            cards = KnowledgeCard.query.filter_by(doc_id=doc_id).all()
        else:
            cards = KnowledgeCard.query.order_by(KnowledgeCard.created_at.desc()).all()
        return jsonify([c.to_dict() for c in cards])

    @app.route('/api/knowledge/cards/<int:card_id>', methods=['GET', 'PUT', 'DELETE'])
    def api_knowledge_card(card_id):
        card = KnowledgeCard.query.get_or_404(card_id)
        if request.method == 'DELETE':
            db.session.delete(card)
            db.session.commit()
            return jsonify({"success": True})
        elif request.method == 'PUT':
            data = request.get_json()
            if 'title' in data:
                card.title = data['title'][:200]
            if 'content' in data:
                card.content = data['content'][:2000]
            if 'tags' in data:
                card.tags = data['tags'][:300]
            db.session.commit()
            return jsonify({"success": True, "card": card.to_dict()})
        return jsonify(card.to_dict())

    @app.route('/api/knowledge/docs/<int:doc_id>', methods=['PUT'])
    def api_knowledge_doc_update(doc_id):
        doc = KnowledgeDoc.query.get_or_404(doc_id)
        data = request.get_json()
        if 'group_name' in data:
            doc.group_name = data['group_name']
        if 'title' in data:
            doc.title = data['title']
        db.session.commit()
        return jsonify({"success": True, "doc": doc.to_dict()})

    # ==================== API: 对标视频 ====================

    @app.route('/api/videos', methods=['GET'])
    def api_videos():
        videos = Video.query.order_by(Video.created_at.desc()).all()
        return jsonify([v.to_dict() for v in videos])

    @app.route('/api/videos/<int:video_id>', methods=['GET', 'DELETE'])
    def api_video(video_id):
        video = Video.query.get_or_404(video_id)
        if request.method == 'DELETE':
            db.session.delete(video)
            db.session.commit()
            return jsonify({"success": True})
        return jsonify(video.to_dict())

    @app.route('/api/videos/parse', methods=['POST'])
    def api_video_parse():
        data = request.get_json()
        url = data.get('url', '')
        if not url:
            return jsonify({"error": "请输入视频链接"}), 400

        try:
            # 调用真实解析器
            parsed = parse_video_url(url)
        except Exception as e:
            return jsonify({"error": f"视频解析失败: {str(e)}"}), 400

        # 构建解析内容摘要（用于AI分析的素材）
        parsed_content_parts = []
        parsed_content_parts.append(f"【视频标题】{parsed.get('title', '未知')}")
        if parsed.get('author'):
            parsed_content_parts.append(f"【作者/账号】{parsed['author']}")
            if parsed.get('author_signature'):
                parsed_content_parts.append(f"【作者简介】{parsed['author_signature'][:200]}")
        if parsed.get('description'):
            parsed_content_parts.append(f"【视频简介】{parsed['description']}")
        if parsed.get('duration'):
            parsed_content_parts.append(f"【时长】{parsed['duration']}")
        if parsed.get('music'):
            parsed_content_parts.append(f"【BGM】{parsed['music']}")
        if parsed.get('tags'):
            parsed_content_parts.append(f"【标签】{', '.join(parsed['tags'])}")

        stats_parts = []
        if parsed.get('play_count'):
            stats_parts.append(f"播放 {parsed['play_count']}")
        if parsed.get('likes'):
            stats_parts.append(f"点赞 {parsed['likes']}")
        if parsed.get('comments'):
            stats_parts.append(f"评论 {parsed['comments']}")
        if parsed.get('shares'):
            stats_parts.append(f"分享 {parsed['shares']}")
        if parsed.get('collects'):
            stats_parts.append(f"收藏 {parsed['collects']}")
        if stats_parts:
            parsed_content_parts.append(f"【数据表现】{' · '.join(stats_parts)}")

        if parsed.get('video_id'):
            parsed_content_parts.append(f"【视频ID】{parsed['video_id']}")
        if parsed.get('real_url'):
            parsed_content_parts.append(f"【真实链接】{parsed['real_url']}")
        if parsed.get('source'):
            parsed_content_parts.append(f"【数据来源】{parsed['source']}")

        has_download = bool(parsed.get('downloaded_file'))
        if has_download:
            parsed_content_parts.append(f"【视频文件】已下载（{parsed.get('file_size', 0) / 1024 / 1024:.1f} MB）")

        if not parsed.get('title') and not parsed.get('description'):
            parsed_content_parts.append("（未能提取到详细内容，可尝试手动输入后再分析）")

        video = Video(
            url=url,
            title=parsed.get('title') or f"视频 {url[-12:]}",
            author=parsed.get('author') or '未知创作者',
            parsed_content='\n'.join(parsed_content_parts),
        )
        db.session.add(video)
        db.session.commit()

        return jsonify({
            "success": True,
            "video": video.to_dict(),
            "parsed_detail": {
                "title": parsed.get('title', ''),
                "author": parsed.get('author', ''),
                "author_signature": parsed.get('author_signature', ''),
                "description": parsed.get('description', ''),
                "tags": parsed.get('tags', []),
                "music": parsed.get('music', ''),
                "duration": parsed.get('duration', ''),
                "cover_url": parsed.get('cover_url', ''),
                "likes": parsed.get('likes', ''),
                "comments": parsed.get('comments', ''),
                "shares": parsed.get('shares', ''),
                "collects": parsed.get('collects', ''),
                "play_count": parsed.get('play_count', ''),
                "video_id": parsed.get('video_id', ''),
                "parse_status": parsed.get('parse_status', 'success'),
                "has_download": has_download,
                "downloaded_file": parsed.get('downloaded_file', ''),
            }
        })

    @app.route('/api/videos/analyze', methods=['POST'])
    def api_video_analyze():
        data = request.get_json()
        api_key = data.get('api_key', '')
        video_id = data.get('video_id')

        video = Video.query.get_or_404(video_id)
        if not api_key:
            return jsonify({"error": "请先设置 DeepSeek API Key"}), 400

        try:
            analysis = analyze_video_content(api_key, video.title, video.parsed_content)
            video.analysis = analysis
            video.is_analyzed = True

            # 自动拆解分析报告为知识卡片
            cards_made = _extract_knowledge_from_analysis(video)

            db.session.commit()
            return jsonify({
                "success": True,
                "video": video.to_dict(),
                "knowledge_cards_created": cards_made,
            })
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route('/api/videos/<int:video_id>/download-video', methods=['GET'])
    def api_video_download(video_id):
        """获取视频下载链接，302重定向到CDN"""
        video = Video.query.get_or_404(video_id)

        from utils.video_parser import parse_video_url as pvu
        try:
            parsed = pvu(video.url)
        except Exception as e:
            return jsonify({"error": f"获取播放地址失败: {str(e)}"}), 500

        play_urls = parsed.get('play_urls', [])
        if not play_urls:
            return jsonify({"error": "未找到可用的视频播放地址"}), 404

        # 返回CDN地址，前端直接打开下载
        return jsonify({
            "success": True,
            "download_url": play_urls[0],
            "title": video.title,
            "note": "链接有效期约2小时，请尽快下载"
        })

    @app.route('/api/videos/<int:video_id>', methods=['PUT'])
    def api_video_update(video_id):
        video = Video.query.get_or_404(video_id)
        data = request.get_json()
        if 'group_name' in data:
            video.group_name = data['group_name']
        if 'title' in data:
            video.title = data['title']
        if 'author' in data:
            video.author = data['author']
        if 'parsed_content' in data:
            video.parsed_content = data['parsed_content']
        db.session.commit()
        return jsonify({"success": True, "video": video.to_dict()})

    @app.route('/api/videos/manual', methods=['POST'])
    def api_video_manual():
        """手动创建视频条目"""
        data = request.get_json()
        title = data.get('title', '').strip()
        if not title:
            return jsonify({"error": "请输入视频标题"}), 400

        video = Video(
            url=data.get('url', '手动输入'),
            title=title,
            author=data.get('author', ''),
            parsed_content=data.get('parsed_content', ''),
        )
        db.session.add(video)
        db.session.commit()
        return jsonify({"success": True, "video": video.to_dict()})

    # ==================== API: 统计 ====================

    @app.route('/api/stats', methods=['GET'])
    def api_stats():
        return jsonify({
            "script_count": Script.query.count(),
            "script_limit": 500,
            "keyword_count": count_keywords(),
            "video_count": Video.query.count(),
        })

    @app.route('/api/scripts/<int:script_id>/download', methods=['GET'])
    def api_script_download(script_id):
        """导出脚本为文档格式 ?format=md|txt|docx"""
        fmt = request.args.get('format', 'md')
        script = Script.query.get_or_404(script_id)
        title = script.title or '脚本'
        req = script.requirement or ''
        content = script.content or ''
        safe = title[:30].replace('/', '_').replace('\\', '_').replace(' ', '_').replace('#', '')

        if fmt == 'md':
            body = f"# {title}\n\n> 创作需求：{req}\n\n---\n\n{content}"
            return body, 200, {'Content-Type': 'text/markdown; charset=utf-8', 'Content-Disposition': f'attachment; filename="{safe}.md"'}

        elif fmt == 'txt':
            sep = '=' * 50
            body = f"{title}\n{sep}\n创作需求：{req}\n{sep}\n\n{content}"
            return body, 200, {'Content-Type': 'text/plain; charset=utf-8', 'Content-Disposition': f'attachment; filename="{safe}.txt"'}

        elif fmt == 'docx':
            try:
                from docx import Document
                from docx.shared import Pt
                from io import BytesIO
                import re

                doc = Document()
                doc.styles['Normal'].font.size = Pt(11)

                doc.add_heading(title, 0)
                if req:
                    doc.add_paragraph(req, style='Intense Quote')

                # 逐行处理：分离表格和非表格
                lines = content.split('\n')
                table_block = []
                for line in lines:
                    s = line.strip()
                    # 收集表格行
                    if s.startswith('|') and s.endswith('|'):
                        table_block.append(s)
                        continue
                    else:
                        # 先输出之前的表格
                        if table_block:
                            _write_table_to_doc(doc, table_block)
                            table_block = []
                    # 输出非表格行
                    if s.startswith('## '):
                        doc.add_heading(s[3:], level=2)
                    elif s.startswith('### '):
                        doc.add_heading(s[4:], level=3)
                    elif s.startswith('# '):
                        doc.add_heading(s[2:], level=1)
                    elif s.startswith('- ') or s.startswith('* '):
                        doc.add_paragraph(s[2:], style='List Bullet')
                    elif re.match(r'^\d+[.)]\s', s):
                        doc.add_paragraph(re.sub(r'^\d+[.)]\s', '', s), style='List Number')
                    elif s == '---':
                        doc.add_paragraph('─' * 40)
                    elif s:
                        p = doc.add_paragraph(s)
                        # 处理粗体标记
                        for run in p.runs:
                            run.text = run.text.replace('**', '')

                # 末尾表格
                if table_block:
                    _write_table_to_doc(doc, table_block)

                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                return buf.getvalue(), 200, {
                    'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'Content-Disposition': f'attachment; filename="{safe}.docx"'
                }
            except ImportError:
                return jsonify({"error": "python-docx 未安装"}), 500

        return jsonify({"error": "不支持的格式"}), 400

    @app.route('/api/admin/seed', methods=['POST'])
    def api_admin_seed():
        """一键导入种子数据到共享数据库"""
        import json as _json
        seed_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'seed_data.json')
        if not os.path.exists(seed_path):
            return jsonify({"error": "种子数据文件不存在"}), 404

        with open(seed_path, 'r', encoding='utf-8') as f:
            data = _json.load(f)

        count = {'plans': 0, 'knowledge_docs': 0, 'knowledge_cards': 0, 'videos': 0, 'scripts': 0}

        for p in data.get('plans', []):
            plan = Plan(title=p.get('title', '')[:200], content=p.get('content', '')[:10000], source_type='import')
            db.session.add(plan)
            count['plans'] += 1

        cards_by_doc = {}
        for card in data.get('knowledge_cards', []):
            doc_id = card.get('doc_id', 'unknown')
            if doc_id not in cards_by_doc:
                cards_by_doc[doc_id] = []
            cards_by_doc[doc_id].append(card)

        for doc_id, cards in cards_by_doc.items():
            doc = KnowledgeDoc(
                title=(cards[0].get('title', '') or '知识文档')[:200],
                file_name='imported.md',
                content='\n\n'.join([c.get('raw_content', c.get('principle', '')) for c in cards[:50]]),
                card_count=len(cards), group_name='导入的知识',
            )
            db.session.add(doc)
            db.session.flush()
            count['knowledge_docs'] += 1
            for card in cards[:50]:
                tags = card.get('tags', '')
                if isinstance(tags, list): tags = ', '.join(tags)
                kc = KnowledgeCard(
                    doc_id=doc.id,
                    title=(card.get('title', '') or '知识卡片')[:200],
                    content=(card.get('raw_content', card.get('principle', '')) or '')[:2000],
                    tags=str(tags)[:300],
                )
                db.session.add(kc)
                count['knowledge_cards'] += 1

        for v in data.get('videos', []):
            video = Video(
                url=v.get('video_url', ''), title=v.get('video_title', '')[:300],
                author=v.get('author_name', '')[:100],
                parsed_content=f"【视频标题】{v.get('video_title','')}\n【作者】{v.get('author_name','')}",
                analysis=v.get('analysis_content', '')[:20000],
                is_analyzed=v.get('analysis_status') == 'analyzed', group_name='全部视频',
            )
            db.session.add(video)
            count['videos'] += 1

        for s in data.get('scripts', []):
            script = Script(
                title=s.get('title', '')[:200], content=s.get('content', '')[:20000],
                requirement=s.get('requirements', '')[:2000],
                keywords=s.get('search_keywords', ''),
            )
            db.session.add(script)
            count['scripts'] += 1

        db.session.commit()
        return jsonify({"success": True, "count": count})

    @app.route('/api/admin/cleanup', methods=['POST'])
    def api_admin_cleanup():
        """去重 + 重建知识卡片"""
        # 1. 视频去重（按URL）
        seen = set()
        dupes = []
        for v in Video.query.order_by(Video.id.desc()).all():
            key = v.url.strip()
            if key in seen:
                dupes.append(v)
            else:
                seen.add(key)
        for d in dupes:
            db.session.delete(d)
        video_dedup = len(dupes)

        # 2. 脚本去重（按标题）
        seen2 = set()
        dupes2 = []
        for s in Script.query.order_by(Script.id.desc()).all():
            key = s.title.strip()
            if key in seen2:
                dupes2.append(s)
            else:
                seen2.add(key)
        for d in dupes2:
            db.session.delete(d)
        script_dedup = len(dupes2)

        # 3. 清除旧知识卡片
        old = KnowledgeCard.query.delete()
        # 清除旧知识文档
        KnowledgeDoc.query.delete()

        # 4. 重新生成知识卡片
        cards_total = 0
        for v in Video.query.filter_by(is_analyzed=True).all():
            if v.analysis and len(v.analysis) > 100:
                n = _extract_knowledge_from_analysis(v)
                cards_total += n
        db.session.commit()

        return jsonify({
            "success": True,
            "videos_removed": video_dedup,
            "scripts_removed": script_dedup,
            "old_cards_removed": old,
            "new_cards_created": cards_total,
            "videos_after": Video.query.count(),
            "scripts_after": Script.query.count(),
            "cards_after": KnowledgeCard.query.count(),
        })

    @app.route('/api/admin/migrate-knowledge', methods=['POST'])
    def api_migrate_knowledge():
        """将已有已分析视频的分析报告拆解为知识卡片"""
        videos = Video.query.filter_by(is_analyzed=True).all()
        total = 0
        for v in videos:
            if v.analysis and len(v.analysis) > 100:
                n = _extract_knowledge_from_analysis(v)
                total += n
        db.session.commit()
        return jsonify({"success": True, "cards_created": total})

    return app


def _write_table_to_doc(doc, rows):
    """将 markdown 表格行列表写入 Word 表格"""
    import re
    data = []
    for r in rows:
        cells = [c.strip() for c in r.split('|')[1:-1]]  # skip leading/trailing empty
        if all(re.match(r'^[-:\s]+$', c) for c in cells):
            continue  # skip separator row
        if cells:
            data.append(cells)
    if not data:
        return
    cols = max(len(r) for r in data)
    table = doc.add_table(rows=len(data), cols=cols)
    table.style = 'Light Grid Accent 1'
    for ri, row in enumerate(data):
        for ci, txt in enumerate(row):
            if ci < cols:
                cell = table.cell(ri, ci)
                cell.text = txt
                if ri == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
    doc.add_paragraph('')


def _smart_card_name(section_title, content):
    """根据分析章节内容生成通俗易懂的卡片名"""
    import re

    # 提取核心关键词
    title_lower = section_title.lower()
    content_first = content[:200].lower()

    # 规则匹配 → 简洁命名
    rules = [
        (r'开头|黄金|钩子|前.*秒|吸引|抓', '钩子'),
        (r'结尾|互动引导|引导|关注|转化|逼单|下单|购买|挂车', '引导转化'),
        (r'文案|台词|话术|口播|对白|金句', '文案话术'),
        (r'爆款|能爆|跑量|流量', '爆款密码'),
        (r'受众|目标|人群|画像|定位', '受众定位'),
        (r'制作|拍摄|运镜|构图|光线|镜头', '拍摄手法'),
        (r'剪辑|节奏|转场|快切|慢镜', '剪辑节奏'),
        (r'bgm|音乐|音效|配乐|卡点', 'BGM音效'),
        (r'标签|话题|关键词|seo', '标签策略'),
        (r'方法论|公式|模板|可复用|底层逻辑', '创作公式'),
        (r'改进|优化|创新|重做', '优化思路'),
        (r'结构|框架|节奏|层次', '结构框架'),
        (r'情绪|情感|共鸣|感动|焦虑', '情绪调动'),
        (r'视觉|画面|字幕|特效|封面', '视觉包装'),
        (r'产品|卖点|植入|种草|展示', '产品植入'),
    ]

    for pattern, name in rules:
        if re.search(pattern, title_lower + content_first):
            # 提取具体的技巧名
            detail = ''
            for line in (section_title + '\n' + content).split('\n'):
                line = line.strip().lstrip('*').strip()
                if len(line) > 5 and len(line) < 50 and line != section_title:
                    detail = line
                    break
            if detail:
                return f"{name}-{detail[:25]}"
            return name

    # 默认：取标题前30字
    short = section_title[:30].strip().lstrip('#').strip()
    return short if short else '创作技巧'


def _extract_knowledge_from_analysis(video):
    """从视频AI分析报告中自动提取知识卡片"""
    analysis = video.analysis
    if not analysis:
        return 0

    sections = analysis.split('\n### ')
    cards_made = 0

    for sec in sections[1:]:
        lines = sec.strip().split('\n')
        raw_title = lines[0].strip()[:200] if lines else ''
        if len(raw_title) < 3:
            continue
        content = '\n'.join(lines[1:]).strip()[:2000] if len(lines) > 1 else sec.strip()[:2000]
        if len(content) < 30:
            continue

        # 智能命名
        title = _smart_card_name(raw_title, content)

        # 标签
        tags = []
        for t in ['钩子', '引导转化', '文案话术', '爆款', '受众', '拍摄', '剪辑', 'BGM',
                   '标签', '公式', '优化', '结构', '情绪', '视觉', '产品']:
            if t in title or t in raw_title or t in content[:100]:
                tags.append(t)

        card = KnowledgeCard(
            doc_id=None,  # 独立卡片，不归属文档
            title=title[:200],
            content=content,
            tags=', '.join(tags[:4]) if tags else '创作技巧',
        )
        db.session.add(card)
        cards_made += 1

    return cards_made


def extract_keywords(text):
    """从文本中提取关键词"""
    common = ['抖音', '短视频', '直播', '带货', '运营', '涨粉', '流量',
              '脚本', '策划', '拍摄', '剪辑', '配乐', '文案', '选题',
              '美妆', '美食', '搞笑', '知识', '育儿', '情感', '干货']
    found = [t for t in common if t in text]
    return ', '.join(found[:5]) if found else '通用'


def count_keywords():
    """统计所有脚本使用过的关键词"""
    scripts = Script.query.all()
    all_keywords = set()
    for s in scripts:
        if s.keywords:
            for k in s.keywords.split(', '):
                all_keywords.add(k)
    return len(all_keywords)


if __name__ == '__main__':
    app = create_app()
    os.makedirs(Config.UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data'), exist_ok=True)
    print("\n" + "=" * 50)
    print("  抖音脚本生成器 V1.0")
    print("  访问地址: http://localhost:5000")
    print("  局域网共享: http://0.0.0.0:5000")
    print("=" * 50 + "\n")
    app.run(host='0.0.0.0', port=5000, debug=True, use_reloader=True)
